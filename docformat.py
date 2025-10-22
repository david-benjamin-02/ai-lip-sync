from docx import Document
import mysql.connector

# Connect to MySQL
conn = mysql.connector.connect(
    host="localhost",
    user="root",
    password="",
    database="hms_doc"
)
cursor = conn.cursor()

# Load DOCX (using a raw string or double backslashes for the file path)
doc = Document(r"C:\Users\prope\Downloads\all_symptoms[1]. aiii.docx")
# doc = Document("C:\\Users\\prope\\Downloads\\all_symptoms.docx")  # Or use double backslashes

# Extract text from the document
text = "\n".join([para.text for para in doc.paragraphs])

# Example logic (you should customize based on how data is structured)
# Loop through each line to find symptom–disease pairs
lines = text.split('\n')
for i in range(len(lines)):
    symptom = lines[i].strip()
    if not symptom or ":" in symptom:  # skip non-symptoms or section headers
        continue
    if i+1 < len(lines):
        diseases = lines[i+1].split(',')
        for disease in diseases:
            disease = disease.strip()
            if disease:
                try:
                    cursor.execute("""
                        INSERT IGNORE INTO symptom_diseases (symptom, disease)
                        VALUES (%s, %s)
                    """, (symptom, disease))
                except Exception as e:
                    print(f"Error inserting: {symptom}, {disease}: {e}")

# Commit the changes and close the connection
conn.commit()
conn.close()
